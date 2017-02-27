import os
import docx

doc = docx.Document()
doc.add_paragraph('Hello world!')
doc.add_paragraph('CTY SUMMER PROGRAM FINAL EVALUATION')



doc.save('helloworld.docx')

names = []

entry = ""

while (entry != "end"):
    first = raw_input("First name?")
    last = raw_input("Last name?")
    names.append(first + " " + last)
    entry = raw_input("end to finish")





