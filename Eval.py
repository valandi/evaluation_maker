import os
import docx
from docx.shared import Pt
import docx.enum.text
from docx.enum.text import WD_LINE_SPACING

doc = docx.Document("C:\Users\Callulis\PycharmProjects\evaluation_maker\Eval_Template.docx")
new_doc = docx.Document()
print doc.paragraphs
fullText = []

style = new_doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

for para in doc.paragraphs:
    font.bold = False
    if "CTY SUMMER" in para.text:
        font.bold = True
        para.alignment = 1
    new_doc.add_paragraph(para.text)

new_doc.save('Eval_Template2.docx')

print fullText
